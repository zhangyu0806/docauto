#!/usr/bin/env python3
"""创建测试数据：PDF发票、Excel数据、Word模板"""

import os
from pathlib import Path

# 确保测试目录存在
test_dir = Path(__file__).parent / "test_data"
test_dir.mkdir(exist_ok=True)


def create_test_pdf():
    """创建包含表格的测试PDF（使用reportlab）"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.units import cm

    output = str(test_dir / "test_invoice.pdf")

    doc = SimpleDocTemplate(output, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()

    # 发票表头
    header_data = [
        ["Invoice Number", "Date", "Company"],
        ["INV-2026-001", "2026-02-26", "Acme Corp"],
    ]
    header_table = Table(header_data, colWidths=[5*cm, 4*cm, 6*cm])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 20))

    # 发票明细
    invoice_data = [
        ["Item", "Quantity", "Unit Price", "Amount"],
        ["Document Processing Service", "100", "0.50", "50.00"],
        ["Data Cleaning Service", "500", "0.20", "100.00"],
        ["Template Generation", "20", "5.00", "100.00"],
        ["OCR Processing", "50", "1.00", "50.00"],
        ["Custom Development", "1", "200.00", "200.00"],
        ["", "", "Total", "500.00"],
    ]
    invoice_table = Table(invoice_data, colWidths=[7*cm, 3*cm, 3*cm, 3*cm])
    invoice_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#f0f0f0')),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
    ]))
    elements.append(invoice_table)

    doc.build(elements)
    print(f"✅ 测试PDF创建: {output}")
    return output


def create_test_pdf_multi_page():
    """创建多页PDF"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Spacer, PageBreak
    from reportlab.lib.units import cm

    output = str(test_dir / "test_multi_page.pdf")
    doc = SimpleDocTemplate(output, pagesize=A4)
    elements = []

    for page in range(1, 4):
        data = [
            ["ID", "Name", "Amount", "Date"],
        ]
        for i in range(1, 11):
            data.append([
                f"{(page-1)*10+i:03d}",
                f"Item-{page}-{i}",
                f"{i * page * 10.5:.2f}",
                f"2026-02-{i:02d}"
            ])

        table = Table(data, colWidths=[3*cm, 5*cm, 4*cm, 4*cm])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ecf0f1')]),
        ]))
        elements.append(table)
        if page < 3:
            elements.append(PageBreak())

    doc.build(elements)
    print(f"✅ 多页测试PDF创建: {output}")
    return output


def create_test_excel():
    """创建需要清洗的测试Excel"""
    import pandas as pd
    import numpy as np

    output = str(test_dir / "test_dirty_data.xlsx")

    data = {
        'Name': ['Alice', 'Bob', '  Charlie  ', 'David', 'Alice', 'Eve', None, 'Frank', 'Grace', 'Bob',
                  'Helen', '  Ivan  ', 'Julia', None, 'Kevin'],
        'Phone': ['13800138000', '138-0013-8001', '13800138002', '+86 138 0013 8003', '13800138000',
                  '138.0013.8004', '', '13800138005', '86-13800138006', '138-0013-8001',
                  '13800138007', '13800138008', '(138)00138009', None, '13800138010'],
        'Date': ['2026-01-15', '2026/02/20', '15-01-2026', '2026.03.10', '2026-01-15',
                 'Jan 5, 2026', None, '2026-04-01', '20260501', '2026/02/20',
                 '2026-06-15', '2026-07-20', '2026/08/25', None, '2026-09-30'],
        'Amount': [1000.5, 2000, 1500.75, None, 1000.5, 3000, None, 500, 750.25, 2000,
                   1200, 800, None, None, 950],
        'Category': ['A', 'B', 'A', 'C', 'A', 'B', None, 'A', 'C', 'B',
                     'A', 'B', 'C', None, 'A'],
    }

    df = pd.DataFrame(data)
    df.to_excel(output, index=False)
    print(f"✅ 测试Excel创建: {output} ({len(df)}行)")
    return output


def create_test_word_template():
    """创建测试Word模板"""
    from docx import Document

    output = str(test_dir / "test_contract_template.docx")

    doc = Document()
    doc.add_heading('Service Agreement', 0)

    doc.add_paragraph('Party A: {{party_a}}')
    doc.add_paragraph('Party B: {{party_b}}')
    doc.add_paragraph('')
    doc.add_paragraph('Project: {{project_name}}')
    doc.add_paragraph('Contract Amount: {{amount}} CNY')
    doc.add_paragraph('Start Date: {{start_date}}')
    doc.add_paragraph('End Date: {{end_date}}')
    doc.add_paragraph('')
    doc.add_paragraph('Terms: {{terms}}')
    doc.add_paragraph('')
    doc.add_paragraph('Signed on: {{sign_date}}')
    doc.add_paragraph('Party A Signature: _______________')
    doc.add_paragraph('Party B Signature: _______________')

    doc.save(output)
    print(f"✅ 测试Word模板创建: {output}")
    return output


def create_test_template_data():
    """创建模板填充数据"""
    import json

    output = str(test_dir / "test_template_data.json")

    data = [
        {
            "party_a": "Acme Corp",
            "party_b": "Tech Solutions Ltd",
            "project_name": "Document Automation System",
            "amount": "50,000",
            "start_date": "2026-03-01",
            "end_date": "2026-06-30",
            "terms": "Monthly payment, 30-day notice for termination",
            "sign_date": "2026-02-26"
        },
        {
            "party_a": "Global Finance Inc",
            "party_b": "Data Processing Co",
            "project_name": "Invoice Processing Service",
            "amount": "30,000",
            "start_date": "2026-04-01",
            "end_date": "2026-09-30",
            "terms": "Quarterly payment, quality guarantee",
            "sign_date": "2026-02-26"
        },
        {
            "party_a": "Legal Partners LLP",
            "party_b": "DocAuto Team",
            "project_name": "Contract Template System",
            "amount": "20,000",
            "start_date": "2026-03-15",
            "end_date": "2026-05-15",
            "terms": "Milestone-based payment",
            "sign_date": "2026-02-26"
        }
    ]

    with open(output, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"✅ 测试模板数据创建: {output} ({len(data)}条)")
    return output


if __name__ == '__main__':
    print("=" * 50)
    print("创建测试数据")
    print("=" * 50)
    create_test_pdf()
    create_test_pdf_multi_page()
    create_test_excel()
    create_test_word_template()
    create_test_template_data()
    print("\n✅ 所有测试数据创建完成！")

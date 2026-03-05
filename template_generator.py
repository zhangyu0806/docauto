#!/usr/bin/env python3
"""
Word模板自动填充脚本
功能：基于模板自动生成文档（合同、报告、信函等）
"""

from docxtpl import DocxTemplate
import pandas as pd
from pathlib import Path
from typing import List, Dict, Any
import json
from datetime import datetime
import re


class WordTemplateGenerator:
    """Word模板生成器"""
    
    def __init__(self, template_path: str):
        """
        初始化生成器
        
        Args:
            template_path: Word模板文件路径（.docx）
        """
        self.template_path = Path(template_path)
        self.template = None
        self.generated_files = []
        
        if not self.template_path.exists():
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        try:
            self.template = DocxTemplate(str(self.template_path))
            print(f"✅ 模板加载成功: {template_path}")
        except Exception as e:
            raise Exception(f"模板加载失败: {e}")
    
    def generate_document(self, data: Dict[str, Any], 
                         output_path: str = None,
                         **context) -> str:
        """
        生成单个文档
        
        Args:
            data: 模板变量字典（如 {'name': '张三', 'amount': 1000}）
            output_path: 输出文件路径（默认：自动生成）
            **context: 额外的上下文变量
            
        Returns:
            生成的文件路径
        """
        # 合并数据
        context.update(data)
        
        # 每次重新加载模板（render后模板状态会改变）
        tpl = DocxTemplate(str(self.template_path))
        
        # 渲染模板
        try:
            tpl.render(context)
            
            # 确定输出路径
            if output_path is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = self.template_path.parent / f"生成_{timestamp}.docx"
            
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 保存文档
            tpl.save(str(output_path))
            print(f"✅ 文档生成成功: {output_path}")
            
            self.generated_files.append(str(output_path))
            return str(output_path)
        
        except Exception as e:
            print(f"❌ 文档生成失败: {e}")
            return None
    
    def batch_generate(self, data_list: List[Dict[str, Any]],
                      output_dir: str = None,
                      naming_pattern: str = "{序号}_{名称}.docx",
                      **global_context) -> List[str]:
        """
        批量生成文档
        
        Args:
            data_list: 数据列表，每项是一个文档的数据字典
            output_dir: 输出目录
            naming_pattern: 文件命名模式，支持变量替换
            **global_context: 全局上下文变量（应用于所有文档）
            
        Returns:
            生成的文件路径列表
        """
        results = []
        
        if output_dir:
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
        else:
            output_path = self.template_path.parent
        
        for idx, data in enumerate(data_list, 1):
            # 添加序号
            data['序号'] = idx
            
            # 合并全局上下文
            context = global_context.copy()
            context.update(data)
            
            # 生成文件名
            try:
                filename = naming_pattern.format(**context)
            except:
                filename = f"文档_{idx}.docx"
            
            # 禁用不安全的文件名字符
            filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
            
            output_file = output_path / filename
            
            # 生成文档
            result = self.generate_document(data, str(output_file), **global_context)
            if result:
                results.append(result)
        
        print(f"\n{'='*50}")
        print(f"批量生成完成: {len(results)}/{len(data_list)}个文档")
        print(f"输出目录: {output_path}")
        print(f"{'='*50}")
        
        return results
    
    def generate_from_excel(self, excel_path: str,
                           sheet_name: str = None,
                           output_dir: str = None,
                           naming_pattern: str = "{序号}_{名称}.docx") -> List[str]:
        """
        从Excel数据批量生成文档
        
        Args:
            excel_path: Excel文件路径
            sheet_name: 工作表名称
            output_dir: 输出目录
            naming_pattern: 文件命名模式
            
        Returns:
            生成的文件路径列表
        """
        # 读取Excel
        df = pd.read_excel(excel_path, sheet_name=sheet_name)
        
        # 转换为字典列表
        data_list = df.to_dict('records')
        
        # 去除NaN值
        for data in data_list:
            for key, value in list(data.items()):
                if pd.isna(value):
                    del data[key]
        
        print(f"从Excel读取了{len(data_list)}条数据")
        
        # 批量生成
        return self.batch_generate(data_list, output_dir, naming_pattern)
    
    def preview_template_variables(self) -> List[str]:
        """
        预览模板中使用的变量
        
        Returns:
            变量名列表
        """
        try:
            variables = self.template.get_undeclared_template_variables()
            return sorted(variables)
        except Exception:
            # Fallback: parse XML manually
            variables = set()
            try:
                doc_xml = self.template.get_xml()
                matches = re.findall(r'\{\{([^}]+)\}\}', doc_xml)
                variables.update(m.strip() for m in matches)
            except:
                pass
            return sorted(variables)


def create_sample_template(output_path: str) -> None:
    """
    创建示例模板（用于测试）
    
    Args:
        output_path: 输出模板路径
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('合同', 0)
    
    # 添加内容
    doc.add_paragraph('甲方：{{甲方名称}}')
    doc.add_paragraph('乙方：{{乙方名称}}')
    
    doc.add_paragraph('根据《中华人民共和国合同法》及相关法律法规，甲乙双方经友好协商，就{{项目名称}}事宜达成如下协议：')
    
    doc.add_heading('一、合同金额', level=1)
    doc.add_paragraph('本合同总金额为人民币{{合同金额}}元（大写：{{金额大写}}）。')
    
    doc.add_heading('二、付款方式', level=1)
    doc.add_paragraph('{{付款方式}}')
    
    doc.add_heading('三、交付时间', level=1)
    doc.add_paragraph('乙方应于{{交付日期}}前完成交付。')
    
    doc.add_heading('四、其他条款', level=1)
    doc.add_paragraph('{{其他条款}}')
    
    # 签名
    doc.add_paragraph('\n\n')
    doc.add_paragraph('甲方（签字）：_________________  日期：{{签订日期}}')
    doc.add_paragraph('乙方（签字）：_________________  日期：{{签订日期}}')
    
    doc.save(output_path)
    print(f"✅ 示例模板已创建: {output_path}")


def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Word模板自动填充工具')
    parser.add_argument('template', help='Word模板文件路径')
    parser.add_argument('-d', '--data', help='JSON数据文件或Excel文件')
    parser.add_argument('-o', '--output', help='输出目录或文件路径')
    parser.add_argument('--sheet', help='Excel工作表名称（如果数据是Excel）')
    parser.add_argument('--pattern', default='{序号}_{名称}.docx', help='批量生成文件命名模式')
    parser.add_argument('--create-sample', action='store_true', help='创建示例模板')
    
    args = parser.parse_args()
    
    # 创建示例模板
    if args.create_sample:
        create_sample_template(args.template)
        return
    
    # 初始化生成器
    try:
        generator = WordTemplateGenerator(args.template)
    except Exception as e:
        print(f"❌ {e}")
        return
    
    # 显示模板变量
    variables = generator.preview_template_variables()
    print(f"\n模板变量: {', '.join(variables)}\n")
    
    # 生成文档
    if args.data:
        data_path = Path(args.data)
        
        if data_path.suffix == '.json':
            # 从JSON加载数据
            with open(data_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            if isinstance(data, list):
                # 批量生成
                generator.batch_generate(data, args.output, args.pattern)
            else:
                # 单个文档
                generator.generate_document(data, args.output)
        
        elif data_path.suffix in ['.xlsx', '.xls']:
            # 从Excel加载数据
            generator.generate_from_excel(
                args.data, 
                args.sheet, 
                args.output, 
                args.pattern
            )
    
    else:
        print("⚠️  请提供数据文件（-d/--data）")
        print("提示：使用 --create-sample 创建示例模板")


if __name__ == '__main__':
    main()
